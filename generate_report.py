"""
generate_report.py — Generate FridgeAI project Word document.

Install deps:
    pip install python-docx

Run:
    python generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
normal_style = doc.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name  = 'Cambria Math'
    run.font.size  = Pt(11)
    run.font.bold  = True
    return p

def add_table_row(table, cells, bold=False, bg=None):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = val
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
        run.font.bold = bold
        run.font.size = Pt(10)
        if bg:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  bg)
            tcPr.append(shd)
    return row


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('FridgeAI')
tr.font.size  = Pt(32)
tr.font.bold  = True
tr.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run('Adaptive Food Waste Reduction System\nTechnical Report')
sr.font.size  = Pt(16)
sr.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(datetime.datetime.now().strftime('%B %Y')).font.size = Pt(12)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — CLEAN DATASETS
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Clean Datasets')
body(
    'Two datasets were used in this project: one for fitting the ASLIE spoilage '
    'probability model, and one for training the visual spoilage classifier.'
)

# ── 1.1 Mendeley ──────────────────────────────────────────────────────────────
h2('1.1  Mendeley Multi-Parameter Fruit Spoilage IoT Dataset')
body(
    'Source: Mendeley Data (publicly available). This dataset was collected from '
    'IoT sensors monitoring fruit storage conditions and provides the ground-truth '
    'binary spoilage labels used to fit the ASLIE logistic regression coefficients.'
)

h3('Dataset Characteristics')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_row(t, ['Property', 'Value'], bold=True, bg='D9E1F2')
rows = [
    ('Total samples',       '10,995 readings'),
    ('Fruit types',         'Banana, Orange, Pineapple, Tomato'),
    ('Label classes',       'Good (fresh) / Bad (spoiled)  —  binary'),
    ('Temperature range',   '21 – 27 °C'),
    ('Humidity range',      '71 – 95 %'),
    ('CO2, Light sensors',  'Excluded (impractical for mixed-fridge use)'),
    ('Class distribution',  'Good ≈ 58%,  Bad ≈ 42%'),
]
for r in rows:
    add_table_row(t, list(r))
doc.add_paragraph()

h3('Preprocessing Steps')
bullet('Features selected: Temperature (°C), Humidity (%), Category encoding (ordinal 1–8).')
bullet('CO2 and Light sensor columns dropped — not available in a standard consumer fridge.')
bullet(
    'Feature normalisation to [0, 1] using fixed reference ranges to ensure numerical '
    'stability across both the warm-storage dataset regime and typical fridge conditions:'
)
code_block(
    'TEMP_NORM  = (0.0,  30.0)   # covers fridge (0-10°C) + dataset (21-27°C)\n'
    'HUMID_NORM = (0.0, 100.0)   # full humidity range\n'
    'CAT_NORM   = (1.0,   8.0)   # ordinal category encoding'
)
bullet('Fruit-to-category mapping applied before training:')
code_block(
    'banana    → fruit    (cat_enc = 5)\n'
    'orange    → fruit    (cat_enc = 5)\n'
    'pineapple → fruit    (cat_enc = 5)\n'
    'tomato    → vegetable (cat_enc = 4)'
)
bullet('Stratified 80/20 train/test split (random_state=42).')
bullet('No duplicate removal required — sensor readings are time-series snapshots.')

h3('Fitted Model Performance')
t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
add_table_row(t2, ['Metric', 'Good (Fresh)', 'Bad (Spoiled)', 'Overall'], bold=True, bg='D9E1F2')
add_table_row(t2, ['Precision',  '0.82', '0.75', '—'])
add_table_row(t2, ['Recall',     '0.76', '0.81', '—'])
add_table_row(t2, ['F1-Score',   '0.79', '0.78', '—'])
add_table_row(t2, ['Accuracy',   '—',    '—',    '0.79'])
add_table_row(t2, ['ROC-AUC',    '—',    '—',    '0.86'])
doc.add_paragraph()

# ── 1.2 Fresh vs Rotten ───────────────────────────────────────────────────────
h2('1.2  Fresh and Rotten Fruits Dataset  (Sriram, Kaggle)')
body(
    'Source: Kaggle — sriramr/fruits-fresh-and-rotten-for-classification. '
    'Used to fine-tune MobileNetV3-Small as a visual spoilage binary classifier.'
)

h3('Dataset Characteristics')
t3 = doc.add_table(rows=1, cols=2)
t3.style = 'Table Grid'
add_table_row(t3, ['Property', 'Value'], bold=True, bg='D9E1F2')
rows3 = [
    ('Total images',      '~13,500'),
    ('Fruit types',       'Apple, Banana, Orange (fresh + rotten variants)'),
    ('Label classes',     'Fresh (0) / Spoiled (1)  —  binary'),
    ('Folder structure',  'train/freshapples, train/rottenapples, … (6 folders)'),
    ('Train split',       '~10,800 images'),
    ('Test split',        '~2,698 images'),
    ('Class balance',     'Fresh 1,164 / Spoiled 1,534 in test set'),
]
for r in rows3:
    add_table_row(t3, list(r))
doc.add_paragraph()

h3('Preprocessing Steps')
bullet('Folder labels remapped to binary: folders prefixed "fresh" → 0, "rotten" → 1.')
bullet('ImageFolder.samples updated (not only .targets) to ensure __getitem__ returns correct binary labels.')
bullet('Images resized to 224 × 224 pixels (MobileNetV3 input requirement).')
bullet('Training augmentations: RandomHorizontalFlip, RandomRotation(±15°), ColorJitter(brightness=0.3, contrast=0.3, saturation=0.2).')
bullet('Normalisation: ImageNet means [0.485, 0.456, 0.406] and stds [0.229, 0.224, 0.225].')
bullet('WeightedRandomSampler applied to handle class imbalance during training.')

h3('Trained Model Performance')
t4 = doc.add_table(rows=1, cols=4)
t4.style = 'Table Grid'
add_table_row(t4, ['Metric', 'Fresh', 'Spoiled', 'Overall'], bold=True, bg='D9E1F2')
add_table_row(t4, ['Precision',  '1.00', '1.00', '—'])
add_table_row(t4, ['Recall',     '1.00', '1.00', '—'])
add_table_row(t4, ['F1-Score',   '1.00', '1.00', '—'])
add_table_row(t4, ['Accuracy',   '—',    '—',    '1.00 (2,698 / 2,698)'])
add_table_row(t4, ['Confusion',  '1164 TN,  0 FP', '0 FN,  1534 TP', '—'])
doc.add_paragraph()

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ARCHITECTURE DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════
h1('2. System Architecture')

body(
    'FridgeAI is a real-time food waste reduction system composed of four layers: '
    'input ingestion, backend inference engine, data persistence, and frontend dashboard. '
    'The diagram below shows the data flow through the system.'
)

h2('2.1  Component Overview')

# Architecture table used as a visual diagram
t5 = doc.add_table(rows=0, cols=1)
t5.style = 'Table Grid'

def arch_row(text, bg='FFFFFF', bold=False, center=False):
    row = t5.add_row()
    cell = row.cells[0]
    cell.text = text
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.bold = bold
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Courier New' if not bold else 'Calibri'
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg)
    tcPr.append(shd)

arch_row('INPUT LAYER', bg='004B87', bold=True, center=True)
arch_row('  [A] Webcam (USB camera)                [B] Barcode scanner         [C] Manual form entry', bg='D6E4F0')
arch_row('       |                                        |                              |',             bg='EBF2FA')
arch_row('  Grounding DINO                     Open Food Facts API              AddItemModal', bg='D6E4F0')
arch_row('  (zero-shot detection)              /lookup/barcode/{code}           POST /items', bg='D6E4F0')
arch_row('  POST /vision/scan                  name + category + shelf_life', bg='D6E4F0')
arch_row('  → DetectedItem[]', bg='D6E4F0')
arch_row('       |', bg='EBF2FA')
arch_row('  MobileNetV3-Small', bg='D6E4F0')
arch_row('  (spoilage classifier)', bg='D6E4F0')
arch_row('  P(spoiled) per bounding box crop', bg='D6E4F0')
arch_row('  spoilage_detected: bool', bg='D6E4F0')

arch_row('', bg='FFFFFF')
arch_row('BACKEND INFERENCE ENGINE  (FastAPI + asyncio)', bg='004B87', bold=True, center=True)
arch_row('  POST /items  →  SQLite INSERT  →  settle_timer.schedule(item_id)', bg='D6E4F0')
arch_row('                                           |', bg='EBF2FA')
arch_row('                               30-min settle delay (asyncio.sleep)', bg='D6E4F0')
arch_row('                                           |', bg='EBF2FA')
arch_row('  ┌────────────────────────────────────────┤', bg='D6E4F0')
arch_row('  │              ASLIE Engine              │', bg='C5D9F1')
arch_row('  │  P_spoil = sigmoid(β₀ + β₁·t          │', bg='C5D9F1')
arch_row('  │            + β₂·T_n + β₃·C_n + β₄·H_n)│', bg='C5D9F1')
arch_row('  │  RSL = binary_search(P_spoil ≥ θ)      │', bg='C5D9F1')
arch_row('  │  RSL = min(RSL_formula, shelf_life - t) │', bg='C5D9F1')
arch_row('  └────────────────────────────────────────┤', bg='D6E4F0')
arch_row('                                           |', bg='EBF2FA')
arch_row('  ┌────────────────────────────────────────┤', bg='D6E4F0')
arch_row('  │              FAPF Engine               │', bg='C5D9F1')
arch_row('  │  S(i) = 0.5·P_spoil + 0.3·Cost_norm   │', bg='C5D9F1')
arch_row('  │         - 0.2·P_consume                │', bg='C5D9F1')
arch_row('  └────────────────────────────────────────┤', bg='D6E4F0')
arch_row('                                           |', bg='EBF2FA')
arch_row('  Alert Engine: P_spoil > 0.80 → CRITICAL', bg='D6E4F0')
arch_row('                P_spoil > 0.50 → WARNING', bg='D6E4F0')
arch_row('                RSL < 1.0 day  → USE_TODAY', bg='D6E4F0')

arch_row('', bg='FFFFFF')
arch_row('REAL-TIME SYNC  (WebSocket broadcast)', bg='004B87', bold=True, center=True)
arch_row('  ITEM_INSERTED  →  frontend ADD_ITEM', bg='D6E4F0')
arch_row('  ITEM_SCORED    →  frontend UPDATE_ITEM  (P_spoil, RSL, fapf_score)', bg='D6E4F0')
arch_row('  ITEM_UPDATED   →  frontend UPDATE_ITEM  (changed_fields)', bg='D6E4F0')
arch_row('  ITEM_DELETED   →  frontend REMOVE_ITEM', bg='D6E4F0')
arch_row('  ALERT_FIRED    →  frontend ADD_ALERT + ADD_TOAST', bg='D6E4F0')

arch_row('', bg='FFFFFF')
arch_row('FRONTEND DASHBOARD  (React + Vite)', bg='004B87', bold=True, center=True)
arch_row('  Inventory view    — live item cards, risk bars, RSL countdown', bg='D6E4F0')
arch_row('  Alerts view       — scrollable alert feed with type badges', bg='D6E4F0')
arch_row('  Analytics view    — FAPF priority table + 7-day spoilage forecast (SVG)', bg='D6E4F0')
arch_row('  ScanModal         — Grounding DINO + barcode (QuaggaJS) camera capture', bg='D6E4F0')
arch_row('  AddItemModal      — manual entry + barcode lookup + shelf-life defaults', bg='D6E4F0')

doc.add_paragraph()

h2('2.2  Shelf-Life Prediction Paths')
t6 = doc.add_table(rows=1, cols=3)
t6.style = 'Table Grid'
add_table_row(t6, ['Path', 'Mechanism', 'Trigger'], bold=True, bg='D9E1F2')
add_table_row(t6, ['Path 1 — ASLIE', 'Logistic regression on time + temp + humidity + category. RSL via binary search.', 'Automatically after 30-min settle delay post item insertion.'])
add_table_row(t6, ['Path 2 — Category defaults', 'Lookup table of default shelf life per ASLIE category (USDA FoodKeeper / FDA).', 'Category dropdown change in AddItemModal.'])
add_table_row(t6, ['Path 3 — Barcode (OFF)', 'Open Food Facts REST API. Maps categories_tags to ASLIE category + default shelf life.', 'Barcode entry in AddItemModal or live scan in ScanModal.'])
doc.add_paragraph()

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PSEUDOCODE + FORMULAS
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Pseudocode and Formulas')

# ── 3.1 ASLIE ─────────────────────────────────────────────────────────────────
h2('3.1  ASLIE — Adaptive Shelf-Life Inference Engine')

body(
    'ASLIE models spoilage probability as a logistic function of elapsed time, '
    'normalised temperature, category encoding, and normalised humidity. Coefficients '
    'β₀, β₂, β₃, β₄ were fitted via logistic regression on the Mendeley dataset. '
    'β₁ was recalibrated analytically so that dairy at 4 °C / 50 % RH reaches '
    'P_spoil = 0.75 in approximately 7 days.'
)

h3('Core Formula')
formula('P_spoil(i, t)  =  σ( β₀  +  β₁·t  +  β₂·T_n  +  β₃·C_n  +  β₄·H_n )')
formula('where  σ(x) = 1 / (1 + e^(−x))')

h3('Feature Normalisation')
formula('T_n  =  (T − T_min) / (T_max − T_min)       [T_min=0, T_max=30 °C]')
formula('H_n  =  (H − H_min) / (H_max − H_min)       [H_min=0, H_max=100 %]')
formula('C_n  =  (C − C_min) / (C_max − C_min)       [C_min=1, C_max=8]')

h3('Fitted Coefficients')
t7 = doc.add_table(rows=1, cols=3)
t7.style = 'Table Grid'
add_table_row(t7, ['Coefficient', 'Value', 'Description'], bold=True, bg='D9E1F2')
add_table_row(t7, ['β₀', '−37.9506', 'Intercept (fitted from Mendeley)'])
add_table_row(t7, ['β₁', '3.40',     'Time decay (days) — recalibrated heuristic'])
add_table_row(t7, ['β₂', '17.0408',  'Temperature (normalised)'])
add_table_row(t7, ['β₃', '−0.0282',  'Category encoding (normalised)'])
add_table_row(t7, ['β₄', '25.9930',  'Humidity (normalised)'])
add_table_row(t7, ['θ',  '0.75',     'Spoilage decision threshold'])
doc.add_paragraph()

h3('Remaining Shelf-Life (RSL) — Binary Search')
body('RSL is the time until P_spoil reaches θ, capped at the declared shelf life:')
formula('RSL  =  min( t* − t_elapsed,   shelf_life − t_elapsed )')
formula('where  t*  =  binary_search { t : P_spoil(t) ≥ θ },   32 iterations')

code_block(
    'FUNCTION rsl(t_elapsed, shelf_life, temp, cat_enc, humidity):\n'
    '    lo ← t_elapsed\n'
    '    hi ← t_elapsed + shelf_life × 3\n'
    '\n'
    '    REPEAT 32 times:\n'
    '        mid ← (lo + hi) / 2\n'
    '        IF P_spoil(mid, temp, cat_enc, humidity) ≥ θ:\n'
    '            hi ← mid\n'
    '        ELSE:\n'
    '            lo ← mid\n'
    '\n'
    '    formula_rsl  ← max(0, lo − t_elapsed)\n'
    '    official_rsl ← max(0, shelf_life − t_elapsed)\n'
    '    RETURN min(formula_rsl, official_rsl)\n'
)

# ── 3.2 FAPF ─────────────────────────────────────────────────────────────────
h2('3.2  FAPF — Freshness-Aware Prioritisation Framework')

body(
    'FAPF assigns each item a priority score S(i) combining spoilage risk, economic '
    'cost, and estimated consumption likelihood. Higher scores indicate items that '
    'should be consumed first.'
)

h3('Scoring Formula')
formula('S(i)  =  0.5 · P_spoil(i)  +  0.3 · Cost_norm(i)  −  0.2 · P_consume(i)')

h3('Term Definitions')
t8 = doc.add_table(rows=1, cols=2)
t8.style = 'Table Grid'
add_table_row(t8, ['Term', 'Definition'], bold=True, bg='D9E1F2')
add_table_row(t8, ['P_spoil(i)', 'Current spoilage probability from ASLIE [0, 1].'])
add_table_row(t8, ['Cost_norm(i)', 'Estimated item cost normalised across all items in inventory [0, 1].'])
add_table_row(t8, ['P_consume(i)', 'Estimated consumption likelihood (currently static; future: learned from user history).'])
doc.add_paragraph()

h3('Pseudocode')
code_block(
    'FUNCTION compute_fapf_scores(items):\n'
    '    max_cost ← max(item.estimated_cost for item in items)\n'
    '\n'
    '    FOR each item in items:\n'
    '        IF max_cost > 0:\n'
    '            cost_norm ← item.estimated_cost / max_cost\n'
    '        ELSE:\n'
    '            cost_norm ← 0\n'
    '\n'
    '        p_consume ← 0.5   // static default (future: user model)\n'
    '\n'
    '        item.fapf_score ← 0.5 × item.P_spoil\n'
    '                        + 0.3 × cost_norm\n'
    '                        − 0.2 × p_consume\n'
    '\n'
    '    RETURN items sorted by fapf_score DESCENDING\n'
)

# ── 3.3 Alert Engine ──────────────────────────────────────────────────────────
h2('3.3  Alert Engine')

h3('Thresholds')
t9 = doc.add_table(rows=1, cols=3)
t9.style = 'Table Grid'
add_table_row(t9, ['Alert Type', 'Condition', 'Action'], bold=True, bg='D9E1F2')
add_table_row(t9, ['CRITICAL',  'P_spoil ≥ 0.80',  'Red toast + alert log entry'])
add_table_row(t9, ['WARNING',   'P_spoil ≥ 0.50',  'Yellow toast + alert log entry'])
add_table_row(t9, ['USE_TODAY', 'RSL < 1.0 day',   'Teal toast + alert log entry'])
doc.add_paragraph()

h3('Pseudocode')
code_block(
    'FUNCTION fire_alerts(item, P_spoil, RSL):\n'
    '    IF P_spoil ≥ ALERT_CRITICAL (0.80):\n'
    '        INSERT alert(type=CRITICAL, item_id, P_spoil, RSL)\n'
    '        BROADCAST ALERT_FIRED event via WebSocket\n'
    '\n'
    '    ELSE IF P_spoil ≥ ALERT_WARNING (0.50):\n'
    '        INSERT alert(type=WARNING, item_id, P_spoil, RSL)\n'
    '        BROADCAST ALERT_FIRED event via WebSocket\n'
    '\n'
    '    IF RSL < ALERT_USE_TODAY (1.0):\n'
    '        INSERT alert(type=USE_TODAY, item_id, P_spoil, RSL)\n'
    '        BROADCAST ALERT_FIRED event via WebSocket\n'
)

# ── 3.4 Visual Spoilage Classifier ───────────────────────────────────────────
h2('3.4  Visual Spoilage Classifier  (MobileNetV3-Small)')

body(
    'A MobileNetV3-Small CNN pre-trained on ImageNet is fine-tuned with a binary '
    'classification head. For each bounding box detected by Grounding DINO, the '
    'crop is passed through this model to obtain P(spoiled).'
)

h3('Architecture Modification')
code_block(
    'MobileNetV3-Small (pretrained, ImageNet)\n'
    '└── Replace classifier head:\n'
    '    Linear(576 → 256)\n'
    '    Hardswish()\n'
    '    Dropout(0.2)\n'
    '    Linear(256 → 1)          ← single logit\n'
    '    sigmoid(logit) → P(spoiled) ∈ [0, 1]\n'
)

h3('Training Configuration')
t10 = doc.add_table(rows=1, cols=2)
t10.style = 'Table Grid'
add_table_row(t10, ['Parameter', 'Value'], bold=True, bg='D9E1F2')
params = [
    ('Loss function',    'BCEWithLogitsLoss'),
    ('Optimiser',        'Adam  (lr = 1e-3)'),
    ('Scheduler',        'CosineAnnealingLR  (T_max = 15)'),
    ('Epochs',           '15'),
    ('Batch size',       '32'),
    ('Image size',       '224 × 224'),
    ('Augmentations',    'RandomHFlip, RandomRotation(15°), ColorJitter'),
    ('Spoilage threshold', '0.5  (P(spoiled) ≥ 0.5 → flagged)'),
]
for p in params:
    add_table_row(t10, list(p))
doc.add_paragraph()

h3('Inference Pseudocode')
code_block(
    'FUNCTION classify_spoilage(bounding_boxes, full_image):\n'
    '    spoilage_scores ← []\n'
    '\n'
    '    FOR each box in bounding_boxes:\n'
    '        crop   ← full_image.crop(box)\n'
    '        tensor ← preprocess(crop)     // resize 224×224, normalise\n'
    '        logit  ← MobileNetV3(tensor)\n'
    '        p_spoiled ← sigmoid(logit)\n'
    '        spoilage_scores.APPEND(p_spoiled)\n'
    '\n'
    '    avg_spoilage ← mean(spoilage_scores)\n'
    '    RETURN avg_spoilage,  (avg_spoilage ≥ 0.5)\n'
)

# ── 3.5 Settle Timer ──────────────────────────────────────────────────────────
h2('3.5  Settle Timer  (Async Scoring Pipeline)')

body(
    'Each item insertion triggers a 30-minute asyncio timer. After the settle delay, '
    'ASLIE and FAPF are computed and results broadcast via WebSocket.'
)

code_block(
    'FUNCTION on_item_inserted(item_id):\n'
    '    schedule asyncio task:\n'
    '        WAIT SETTLE_DELAY (1800 seconds)\n'
    '        t_elapsed ← now() − item.entry_time  (days)\n'
    '        P_spoil, RSL ← ASLIE.compute(t_elapsed, temp, shelf_life, cat_enc, humidity)\n'
    '        fapf_score   ← FAPF.score(P_spoil, cost_norm)\n'
    '        UPDATE items SET P_spoil, RSL, fapf_score, confidence_tier WHERE item_id\n'
    '        BROADCAST ITEM_SCORED event\n'
    '        fire_alerts(item, P_spoil, RSL)\n'
    '\n'
    'ON startup:\n'
    '    FOR each item in database WHERE P_spoil IS NULL:\n'
    '        reschedule settle timer\n'
)

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = 'FridgeAI_Technical_Report.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
